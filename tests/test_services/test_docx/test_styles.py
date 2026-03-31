import io
from docx import Document
from docx.shared import Pt, Inches, RGBColor

from app.services.docx.styles import (
    HEADING_FONT, BODY_FONT, NAVY, WHITE, BODY_TEXT_COLOR,
    SIDEBAR_BG, ROLE_COLOR, PROFILE_COLOR,
    NAME_SIZE, TITLE_SIZE, SECTION_HEADER_SIZE, BODY_SIZE,
    SMALL_SIZE, CONTACT_SIZE,
    PAGE_WIDTH, PAGE_HEIGHT,
    MARGIN_TOP, MARGIN_RIGHT, MARGIN_BOTTOM, MARGIN_LEFT,
    SIDEBAR_WIDTH, MAIN_WIDTH,
    set_run_font, setup_page, add_header_band,
    add_section_header, create_two_column_table,
    add_contact_items, add_no_border_table,
)


def test_color_constants():
    assert NAVY == RGBColor(0x23, 0x4F, 0x77)
    assert WHITE == RGBColor(0xFF, 0xFF, 0xFF)
    assert BODY_TEXT_COLOR == RGBColor(0x2D, 0x26, 0x2E)
    assert SIDEBAR_BG == "629DD1"
    assert ROLE_COLOR == RGBColor(0x3B, 0x38, 0x38)
    assert PROFILE_COLOR == RGBColor(0x3B, 0x38, 0x38)


def test_font_constants():
    assert HEADING_FONT == "Century Gothic"
    assert BODY_FONT == "Calibri Light"


def test_size_constants():
    assert NAME_SIZE == Pt(28)
    assert TITLE_SIZE == Pt(14)
    assert SECTION_HEADER_SIZE == Pt(15)
    assert BODY_SIZE == Pt(10)
    assert SMALL_SIZE == Pt(10)
    assert CONTACT_SIZE == Pt(9)


def test_page_dimensions():
    assert PAGE_WIDTH == 11906
    assert PAGE_HEIGHT == 16838
    assert MARGIN_TOP == 630
    assert MARGIN_RIGHT == 282
    assert MARGIN_BOTTOM == 1440
    assert MARGIN_LEFT == 1440


def test_column_widths():
    assert SIDEBAR_WIDTH == PAGE_WIDTH * 2 // 5
    assert MAIN_WIDTH == PAGE_WIDTH - SIDEBAR_WIDTH


def test_setup_page():
    doc = Document()
    setup_page(doc)
    section = doc.sections[0]
    from docx.shared import Emu
    assert section.page_width == Emu(PAGE_WIDTH * 914400 // 1440)
    assert section.page_height == Emu(PAGE_HEIGHT * 914400 // 1440)


def test_add_header_band():
    doc = Document()
    add_header_band(doc, "John Doe", "Engineer")
    assert len(doc.tables) == 1
    cell = doc.tables[0].cell(0, 0)
    assert "JOHN DOE" in cell.text
    assert "Engineer" in cell.text


def test_create_two_column_table():
    doc = Document()
    sidebar, main = create_two_column_table(doc)
    assert len(doc.tables) == 1
    assert sidebar is not None
    assert main is not None


def test_add_section_header():
    doc = Document()
    sidebar, _ = create_two_column_table(doc)
    add_section_header(sidebar, "CONTACT")
    assert "CONTACT" in sidebar.text


def test_add_contact_items():
    doc = Document()
    sidebar, _ = create_two_column_table(doc)
    contact = {"email": "test@test.com", "phone": "123", "location": "NYC"}
    add_contact_items(sidebar, contact)
    assert "test@test.com" in sidebar.text
    assert "123" in sidebar.text
    assert "NYC" in sidebar.text
