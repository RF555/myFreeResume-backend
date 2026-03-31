import io
from docx import Document

from app.services.docx.resume_generator import generate_resume
from app.models.entry import (
    ResumeData, Contact, ExperienceItem, EducationItem, SkillCategory,
)


def _all_text(doc):
    text = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += "\n" + cell.text
    return text


def test_generate_resume_returns_valid_docx():
    data = ResumeData(
        name="John Doe",
        professional_title="Software Engineer",
        summary="Experienced developer.",
        contact=Contact(email="john@example.com", phone="+1234567890", location="New York"),
        skill_highlights=["Problem Solving", "Communication"],
        skill_categories=[
            SkillCategory(name="Advanced", skills=["Python", "React"]),
            SkillCategory(name="Intermediate", skills=["Go"]),
        ],
        experience=[
            ExperienceItem(
                company="Google", role="Senior Engineer",
                start_date="2020", end_date="Present",
                bullets=["Led team of 5", "Built microservices"],
            )
        ],
        education=[
            EducationItem(institution="MIT", degree="B.S. CS", start_date="2014", end_date="2018")
        ],
        languages=["English", "Spanish"],
    )
    result = generate_resume(data)
    assert isinstance(result, bytes)
    doc = Document(io.BytesIO(result))
    text = _all_text(doc)
    assert "JOHN DOE" in text
    assert "Software Engineer" in text


def test_resume_has_two_column_layout():
    data = ResumeData(name="Jane", professional_title="Dev")
    doc = Document(io.BytesIO(generate_resume(data)))
    assert len(doc.tables) >= 2
    body_table = doc.tables[1]
    assert len(body_table.columns) == 2


def test_resume_sidebar_contains_contact():
    data = ResumeData(
        name="Jane",
        professional_title="Dev",
        contact=Contact(email="jane@test.com", phone="555-1234"),
    )
    doc = Document(io.BytesIO(generate_resume(data)))
    sidebar = doc.tables[1].cell(0, 0)
    assert "jane@test.com" in sidebar.text
    assert "555-1234" in sidebar.text


def test_resume_sidebar_contains_education():
    data = ResumeData(
        name="Jane",
        professional_title="Dev",
        education=[EducationItem(degree="B.S. CS", institution="MIT", end_date="2020")],
    )
    doc = Document(io.BytesIO(generate_resume(data)))
    sidebar = doc.tables[1].cell(0, 0)
    assert "B.S. CS" in sidebar.text
    assert "MIT" in sidebar.text


def test_resume_sidebar_contains_skill_categories():
    data = ResumeData(
        name="Jane",
        professional_title="Dev",
        skill_categories=[SkillCategory(name="Advanced", skills=["Python", "React"])],
    )
    doc = Document(io.BytesIO(generate_resume(data)))
    sidebar = doc.tables[1].cell(0, 0)
    assert "Advanced" in sidebar.text
    assert "Python" in sidebar.text


def test_resume_sidebar_contains_languages_in_skill_highlights():
    data = ResumeData(
        name="Jane",
        professional_title="Dev",
        languages=["English", "Hebrew"],
        skill_highlights=["Leadership"],
    )
    doc = Document(io.BytesIO(generate_resume(data)))
    sidebar = doc.tables[1].cell(0, 0)
    assert "English" in sidebar.text
    assert "Hebrew" in sidebar.text
    assert "Leadership" in sidebar.text


def test_resume_main_contains_experience():
    data = ResumeData(
        name="Jane",
        professional_title="Dev",
        experience=[ExperienceItem(role="Engineer", company="Acme", start_date="2020", end_date="2022")],
    )
    doc = Document(io.BytesIO(generate_resume(data)))
    main = doc.tables[1].cell(0, 1)
    assert "Engineer" in main.text
    assert "Acme" in main.text


def test_resume_main_contains_summary():
    data = ResumeData(name="Jane", professional_title="Dev", summary="Great dev.")
    doc = Document(io.BytesIO(generate_resume(data)))
    main = doc.tables[1].cell(0, 1)
    assert "Great dev." in main.text


def test_resume_respects_hidden_sections():
    data = ResumeData(
        name="Jane",
        professional_title="Dev",
        summary="Great dev.",
        skill_highlights=["Leadership"],
    )
    result = generate_resume(data, hidden_sections={"summary": True, "skill_highlights": True})
    doc = Document(io.BytesIO(result))
    main = doc.tables[1].cell(0, 1)
    sidebar = doc.tables[1].cell(0, 0)
    assert "Great dev." not in main.text
    assert "Leadership" not in sidebar.text


def test_resume_empty_data():
    data = ResumeData()
    result = generate_resume(data)
    assert isinstance(result, bytes)
