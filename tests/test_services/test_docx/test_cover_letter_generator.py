import io
from docx import Document

from app.services.docx.cover_letter_generator import generate_cover_letter
from app.models.entry import CoverLetterData, Contact


def _all_text(doc):
    text = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += "\n" + cell.text
    return text


def test_cover_letter_returns_valid_docx():
    data = CoverLetterData(
        addressee_name="Jane Smith",
        addressee_company="Google",
        date="March 25, 2026",
        salutation="Dear Ms. Smith,",
        body_paragraphs=["I am writing to express my interest.", "I look forward to hearing from you."],
        closing="Best regards",
    )
    contact = Contact(email="john@example.com", phone="+1234567890", location="New York")
    result = generate_cover_letter("John Doe", "Software Engineer", contact, data)
    assert isinstance(result, bytes)
    doc = Document(io.BytesIO(result))
    text = _all_text(doc)
    assert "JOHN DOE" in text


def test_cover_letter_has_two_column_layout():
    data = CoverLetterData()
    contact = Contact()
    doc = Document(io.BytesIO(generate_cover_letter("Jane", "Dev", contact, data)))
    assert len(doc.tables) >= 2
    body_table = doc.tables[1]
    assert len(body_table.columns) == 2


def test_cover_letter_sidebar_contains_contact():
    data = CoverLetterData()
    contact = Contact(email="jane@test.com", phone="555")
    doc = Document(io.BytesIO(generate_cover_letter("Jane", "Dev", contact, data)))
    sidebar = doc.tables[1].cell(0, 0)
    assert "jane@test.com" in sidebar.text


def test_cover_letter_sidebar_contains_addressee():
    data = CoverLetterData(addressee_name="Bob Jones", addressee_company="Acme Corp")
    contact = Contact()
    doc = Document(io.BytesIO(generate_cover_letter("Alice", "Dev", contact, data)))
    sidebar = doc.tables[1].cell(0, 0)
    assert "Bob Jones" in sidebar.text
    assert "Acme Corp" in sidebar.text


def test_cover_letter_main_contains_salutation():
    data = CoverLetterData(salutation="Dear Mr. Smith,")
    contact = Contact()
    doc = Document(io.BytesIO(generate_cover_letter("Alice", "Dev", contact, data)))
    main = doc.tables[1].cell(0, 1)
    assert "DEAR MR. SMITH," in main.text


def test_cover_letter_main_contains_body():
    data = CoverLetterData(
        body_paragraphs=["First paragraph.", "Second paragraph."],
        closing="Sincerely",
    )
    contact = Contact()
    doc = Document(io.BytesIO(generate_cover_letter("Alice", "Dev", contact, data)))
    main = doc.tables[1].cell(0, 1)
    assert "First paragraph." in main.text
    assert "Second paragraph." in main.text
    assert "Sincerely" in main.text
    assert "Alice" in main.text


def test_cover_letter_empty_data():
    data = CoverLetterData()
    contact = Contact()
    result = generate_cover_letter("", "", contact, data)
    assert isinstance(result, bytes)
